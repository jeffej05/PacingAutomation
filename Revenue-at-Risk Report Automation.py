#!/usr/bin/env python
# coding: utf-8

# In[3]:


#import packages
import smtplib
import glob
import docx
import pyodbc
import os 
import pandas as pd
import openpyxl as op
import re
import numpy as np
import seaborn as sns
import sys
import time
from datetime import datetime
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from email.mime.base import MIMEBase
from email.encoders import encode_base64
from email import encoders
import gspread
import pprint
from oauth2client.service_account import ServiceAccountCredentials


# In[4]:


#Current Month Variable for Pivot Column Name Change
current_month_text = datetime.now().strftime('%B')
current_month_text


# In[5]:


#Current Month Variable for Pivot Word Doc labels
current_month_year_text = datetime.now().strftime('%B %Y')
current_month_year_text


# In[6]:


#Use current path were import file is stored
os.chdir("/Users/Jeffrey.Lu@ibm.com/Desktop/")


# In[7]:


# When importing data, save the file to your cwd with the below formatted titel
df=pd.read_csv(time.strftime('Pacing Report %m-%d-%Y'))


# In[8]:


#Removing Unnecessary Columns
df.drop(['Ad_Computed_Status', 'Month_Goal_Amount','Month_Captured_Amount',        'Total_Captured_Amount','Current_Month_Overdelivery_Amount','Total_Overdelivery_Amount'],       axis = 1, inplace = True)


# In[9]:


#Convert UD Column DType to float
df['Current_Month_Estimated_Underdelivery_Amount'] = df['Current_Month_Estimated_Underdelivery_Amount'].str.replace(',','').astype(float)

df['Estimated_Underdelivery_Amount'] = df['Estimated_Underdelivery_Amount'].str.replace(',','').astype(float)


# In[10]:


# Printing and Dropping Bad AV lines

Bad_AV_Lines = df['Ad_Name'][(df['Ad_Name'].str.contains('_AV', regex = True))&              (df['Current_Month_Estimated_Underdelivery_Amount']!= 0)]

if df['Ad_Name'][(df['Ad_Name'].str.contains('_AV', regex = True))&              (df['Current_Month_Estimated_Underdelivery_Amount']!= 0)].count() > 0:
    Bad_AV_IndexNames = Bad_AV_Lines.index
    df.drop(Bad_AV_IndexNames , inplace=True)
    print('The following lines have been dropped: ' + Bad_AV_Lines)
else: 
    print('_AV lines are good, none show underdelivery')
    


# In[11]:


# Priting and Dropping Bad MG lines

Bad_MG_Lines = df['Ad_Name'][(df['Ad_Name'].str.contains('_MG', regex = True))&              (df['Current_Month_Estimated_Underdelivery_Amount']!= 0)]

if df['Ad_Name'][(df['Ad_Name'].str.contains('_MG', regex = True))&              (df['Current_Month_Estimated_Underdelivery_Amount']!= 0)].count() > 0:
    Bad_MG_IndexNames = Bad_MG_Lines.index
    df.drop(Bad_MG_IndexNames , inplace=True)
    print('The following lines have been dropped: ' + Bad_MG_Lines)
else: 
    print('_MG lines are good, none show underdelivery')


# In[12]:


#Monday Report Pivot Table Total
#No try/except or if statement on this line. We want the program to crash if pivot_total is 0 or empty
df_pivot_total= pd.pivot_table(df, index= 'Advertiser', values= 'Current_Month_Estimated_Underdelivery_Amount', aggfunc= np.sum)
df_pivot_total= df_pivot_total[df_pivot_total.values != 0].sort_values('Current_Month_Estimated_Underdelivery_Amount', ascending = False)
df_pivot_total


# In[13]:


#Completed Pivot; The try/except is to adjust for pivots that don't have values as the 7th line will error out
#if statement is to account for zero-contained lines that were taken out of the pivot, leaving an empty DF object. 
df_pivot_complete = pd.pivot_table(df.where(df['Ad_Status']=='complete'), index= 'Advertiser', values = 'Current_Month_Estimated_Underdelivery_Amount', aggfunc=np.sum)

try:
    df_pivot_complete= df_pivot_complete[df_pivot_complete.values != 0].sort_values('Current_Month_Estimated_Underdelivery_Amount', ascending = False)
    print(df_pivot_complete)
except:
    df_pivot_complete = '$0 Completed Lines'
    print(df_pivot_complete)
    
#Regex Objects    
ZeroRegex=re.compile(r'\$0')
moComplete=ZeroRegex.findall(str(df_pivot_complete))
moComplete

#Forces empty dataframes to take on string of $0 lines if it did not take on already. reason: empty frames won't take on '$0'.

try:
    if (type(df_pivot_complete) != str) & (df_pivot_complete.empty == True):
        df_pivot_complete = '$0 Completed Lines'
        print(df_pivot_complete)
    elif (moComplete == ['$0']) & (df_pivot_complete != '$0 Completed Lines'):
        df_pivot_complete = '$0 Completed Lines'
        print(df_pivot_complete)
    elif (type(df_pivot_complete) == str) & (df_pivot_complete != '$0 Completed Lines'):
        df_pivot_suspended = '$0 Completed Lines'
        print(df_pivot_suspended)
except:
    print(df_pivot_complete)


# In[14]:


Completed = df['Current_Month_Estimated_Underdelivery_Amount'].where(df['Ad_Status']=='complete').fillna(0).sum()


# In[15]:


Completed


# In[16]:


#Suspended Pivot
df_pivot_suspended = pd.pivot_table(df.where(df['Ad_Status']=='suspended'), index= 'Advertiser', values = 'Current_Month_Estimated_Underdelivery_Amount', aggfunc=np.sum)

try:
    df_pivot_suspended= df_pivot_suspended[df_pivot_suspended.values != 0].sort_values('Current_Month_Estimated_Underdelivery_Amount', ascending = False)
    print(df_pivot_suspended)
except:
    df_pivot_suspended = '$0 Suspended Lines'
    print(df_pivot_suspended)

#REGEX OBJECTS
ZeroRegex=re.compile(r'\$0')
moSuspended=ZeroRegex.findall(str(df_pivot_suspended))
moSuspended

#Forces empty dataframes to take on string of $0 lines if it did not take on already. reason: empty frames won't take on '$0'.
try:
    if (type(df_pivot_suspended) != str) & (df_pivot_suspended.empty == True):
        df_pivot_suspended = '$0 Suspended Lines'
        print(df_pivot_suspended)
    elif (moSuspended == ['$0']) & (df_pivot_suspended != '$0 Suspended Lines'):
        df_pivot_suspended = '$0 Suspended Lines'
        print(df_pivot_suspended)
    elif (type(df_pivot_suspended) == str) & (df_pivot_suspended != '$0 Suspended Lines'):
        df_pivot_suspended = '$0 Suspended Lines'
        print(df_pivot_suspended)

except:
    print(df_pivot_suspended)

###Elementwise comparison failure warning needs python 3.73 to be resolved. Make sure to use correct version of python.


# In[17]:


#Converting Pivot Table Results to calculatable variable
Suspended = df['Current_Month_Estimated_Underdelivery_Amount'].where(df['Ad_Status']=='suspended').fillna(0).sum()


# In[18]:


Suspended


# In[19]:


#Sponsorship Pivot
df_pivot_sponsorship = pd.pivot_table(df.where((df['Ad_Status']=='active')&(df['Ad_Type']=='SPONSORSHIP')), index= 'Advertiser', values = 'Current_Month_Estimated_Underdelivery_Amount', aggfunc=np.sum)

try:
    df_pivot_sponsorship= df_pivot_sponsorship[df_pivot_sponsorship.values != 0].sort_values('Current_Month_Estimated_Underdelivery_Amount', ascending = False)
    print(df_pivot_sponsorship)
except:
    df_pivot_sponsorship='$0 Sponsorship Lines'
    print(df_pivot_sponsorship)

#REGEX OBJECTS  
ZeroRegex=re.compile(r'\$0')
moSponsorship=ZeroRegex.findall(str(df_pivot_sponsorship))
moSponsorship

#Forces empty dataframes to take on string of $0 lines

try:
    if (type(df_pivot_sponsorship) != str) & (df_pivot_sponsorship.empty == True):
        df_pivot_sponsorship = '$0 Sponsorship Lines'
        print(df_pivot_sponsorship)
    elif (moSponsorship == ['$0']) & (df_pivot_sponsorship != '$0 Sponsorship Lines'):
        df_pivot_sponsorship = '$0 Sponsorship Lines'
        print(df_pivot_sponsorship)
    elif (type(df_pivot_sponsorship) == str) & (df_pivot_sponsorship != '$0 Sponsorship Lines'):
        df_pivot_sponsorship = '$0 Sponsorship Lines'
        print(df_pivot_sponsorship)

except:
    print(df_pivot_sponsorship)


# In[20]:


#WFX Pivot
df_pivot_WFX = pd.pivot_table(df[~df['Product_Name'].str.contains('Preroll', na = False)]                              [df['Ad_Type'].str.contains('STANDARD|HOUSE|AD_EXCHANGE', na = False)]                              .where((df['Ad_Status']=='active')&(df['Product_Category']=='WEATHRFX')),                              index= 'Advertiser',values = 'Current_Month_Estimated_Underdelivery_Amount',                              aggfunc=np.sum)

try:
    df_pivot_WFX= df_pivot_WFX[df_pivot_WFX.values != 0].sort_values('Current_Month_Estimated_Underdelivery_Amount', ascending = False)
    print(df_pivot_WFX)
except:
    df_pivot_WFX='$0 WFX Lines'
    print(df_pivot_WFX)
    
#REGEX OBJECTS  
ZeroRegex=re.compile(r'\$0')
moWFX=ZeroRegex.findall(str(df_pivot_WFX))
moWFX

#Forces empty dataframes to take on string of $0 lines

try:
    if (type(df_pivot_WFX) != str) & (df_pivot_WFX.empty == True):
        df_pivot_WFX = '$0 WFX Lines'
        print(df_pivot_WFX)
    elif (moWFX == ['$0']) & (df_pivot_WFX != '$0 WFX Lines'):
        df_pivot_WFX = '$0 WFX Lines'
        print(df_pivot_WFX)
    elif (type(df_pivot_WFX) == str) & (df_pivot_WFX != '$0 WFX Lines'):
        df_pivot_sponsorship = '$0 WFX Lines'
        print(df_pivot_WFX)
except:
    print(df_pivot_WFX)


# In[21]:


#AFX Pivot
df_pivot_AFX = pd.pivot_table(df[~df['Product_Name'].str.contains('Preroll', na = False)]                              [df['Ad_Type'].str.contains('STANDARD|HOUSE|AD_EXCHANGE', na = False)]                              .where((df['Ad_Status']=='active')&(df['Product_Category']=='AUDIENCEFX')),                              index= 'Advertiser',values = 'Current_Month_Estimated_Underdelivery_Amount',                              aggfunc=np.sum)

try:
    df_pivot_AFX= df_pivot_AFX[df_pivot_AFX.values != 0].sort_values('Current_Month_Estimated_Underdelivery_Amount', ascending = False)
    print(df_pivot_AFX)
except:
    df_pivot_WFX='$0 AFX Lines'
    print(df_pivot_AFX)
    
#REGEX OBJECTS  
ZeroRegex=re.compile(r'\$0')
moAFX=ZeroRegex.findall(str(df_pivot_AFX))
moAFX

#Forces empty dataframes to take on string of $0 lines

try:
    if (type(df_pivot_AFX) != str) & (df_pivot_AFX.empty == True):
        df_pivot_AFX = '$0 AFX Lines'
        print(df_pivot_AFX)
    elif (moAFX == ['$0']) & (df_pivot_AFX != '$0 AFX Lines'):
        df_pivot_AFX = '$0 AFX Lines'
        print(df_pivot_AFX)
    elif (type(df_pivot_AFX) == str) & (df_pivot_AFX != '$0 AFX Lines'):
        df_pivot_sponsorship = '$0 AFX Lines'
        print(df_pivot_AFX)
except:
    print(df_pivot_AFX)              
                                                                       


# In[22]:


#JFX Pivot
df_pivot_JFX = pd.pivot_table(df[~df['Product_Name'].str.contains('Preroll', na = False)]                              [df['Ad_Type'].str.contains('STANDARD|HOUSE|AD_EXCHANGE', na = False)]                              .where((df['Ad_Status']=='active')&(df['Product_Category']=='JFX')),                              index= 'Advertiser',values = 'Current_Month_Estimated_Underdelivery_Amount',                              aggfunc=np.sum)

try:
    df_pivot_JFX= df_pivot_JFX[df_pivot_AFX.values != 0].sort_values('Current_Month_Estimated_Underdelivery_Amount', ascending = False)
    print(df_pivot_JFX)
except:
    df_pivot_JFX='$0 JFX Lines'
    print(df_pivot_JFX)
    
#REGEX OBJECTS  
ZeroRegex=re.compile(r'\$0')
moJFX=ZeroRegex.findall(str(df_pivot_JFX))
moJFX

#Forces empty dataframes to take on string of $0 lines

try:
    if (type(df_pivot_JFX) != str) & (df_pivot_JFX.empty == True):
        df_pivot_JFX = '$0 JFX Lines'
        print(df_pivot_JFX)
    elif (moJFX == ['$0']) & (df_pivot_JFX != '$0 JFX Lines'):
        df_pivot_JFX = '$0 JFX Lines'
        print(df_pivot_JFX)
    elif (type(df_pivot_JFX) == str) & (df_pivot_JFX != '$0 JFX Lines'):
        df_pivot_sponsorship = '$0 JFX Lines'
        print(df_pivot_JFX)
except:
    print(df_pivot_JFX)   


# In[23]:


#Preroll Pivot
df_pivot_preroll = pd.pivot_table(df[df['Product_Name'].str.contains('Preroll', na = False)]                              [df['Ad_Type'].str.contains('STANDARD|HOUSE|AD_EXCHANGE', na = False)]                              .where((df['Ad_Status']=='active')),                              index= 'Advertiser',values = 'Current_Month_Estimated_Underdelivery_Amount',                              aggfunc=np.sum)

try:
    df_pivot_preroll= df_pivot_preroll[df_pivot_preroll.values != 0].sort_values('Current_Month_Estimated_Underdelivery_Amount', ascending = False)
    print(df_pivot_preroll)
except:
    df_pivot_preroll='$0 Preroll Lines'
    print(df_pivot_preroll)
    
    
#REGEX OBJECTS  
ZeroRegex=re.compile(r'\$0')
moPreroll=ZeroRegex.findall(str(df_pivot_preroll))
moPreroll

#Forces empty dataframes to take on string of $0 lines

try:
    if (type(df_pivot_preroll) != str) & (df_pivot_preroll.empty == True):
        df_pivot_preroll = '$0 Preroll Lines'
        print(df_pivot_preroll)
    elif (moPreroll == ['$0']) & (df_pivot_preroll != '$0 Preroll Lines'):
        df_pivot_preroll = '$0 Preroll Lines'
        print(df_pivot_preroll)
    elif (type(df_pivot_preroll) == str) & (df_pivot_preroll != '$0 Preroll Lines'):
        df_pivot_preroll = '$0 Preroll Lines'
        print(df_pivot_preroll)

except:
    print(df_pivot_preroll)    


    
#This warning is based off the differences in opinions between developers. If methodology changes in the future, code may need to be rewritten


# In[24]:


#Advertiser Pivot
df_pivot_advertiser= pd.pivot_table(df, index= 'Advertiser', values= ['Current_Month_Estimated_Underdelivery_Amount','Estimated_Underdelivery_Amount'], aggfunc= np.sum)
df_pivot_advertiser= df_pivot_advertiser[df_pivot_advertiser.values != [0,0]].sort_values('Current_Month_Estimated_Underdelivery_Amount', ascending = False)
df_pivot_advertiser = df_pivot_advertiser.drop_duplicates()
df_pivot_advertiser['Notes for Campaigns with $10k + Risk for the Current Month']=("")
df_pivot_advertiser.rename(columns={'Current_Month_Estimated_Underdelivery_Amount': current_month_text +' Risk'}, inplace=True)
df_pivot_advertiser.rename(columns={'Estimated_Underdelivery_Amount': 'Total Risk'}, inplace=True)


df_pivot_advertiser


# In[25]:


#Changing Pivot column title to have dynamic month label
df_pivot_advertiser[current_month_text + ' Risk']=df_pivot_advertiser[current_month_text + ' Risk'].apply(lambda x: '${:,.2f}'.format(x))


# In[26]:


#Incorporating currency formatting
df_pivot_advertiser['Total Risk']=df_pivot_advertiser['Total Risk'].apply(lambda x: '${:,.2f}'.format(x))


# In[27]:


df_pivot_advertiser


# In[28]:


##### Late Creative Section #####


# In[29]:


#Late Creative file should be saved to cwd with the below formatting structure
df2=pd.read_csv(time.strftime('Late Creative %m-%d-%Y'), skipinitialspace=True)


# In[30]:


df2.head()


# In[31]:


#Deleting UK Lines
df2= df2[~df2.Region.isin(['UK', 'United Kingdom', '-'])]


# In[32]:


#Deleting Bidopt Lines
df2= df2[~df2.Name.str.contains('bidopt')]


# In[33]:


df2.reset_index(inplace = True)


# In[34]:


#Stripping CSV whitespace in column
df2['Revenue_Impact']=df2['Revenue_Impact'].str.strip()


# In[35]:


# Converting Money series to float series
df2['Revenue_Impact'] = df2['Revenue_Impact'].replace( '[\$,)]','', regex=True ).astype(float)


# In[36]:


#Stripping CSV whitespace in column
df2['Total_Month_Dollars']=df2['Total_Month_Dollars'].str.strip()
df2['Total_Month_Dollars'] = df2['Total_Month_Dollars'].replace( '[\$,)]','', regex=True ).astype(float)


# In[38]:


df2['Total_Month_Dollars'][0]


# In[39]:


#Late Creative Pivot
late_creative_pivot = pd.pivot_table(df2, index= ['Account_Executive', 'Advertiser'],values = ['Name','Total_Month_Dollars','Revenue_Impact'],aggfunc = {'Name': pd.Series.nunique, 'Total_Month_Dollars': np.sum, 'Revenue_Impact':np.sum}, margins = True)
late_creative_pivot


# In[41]:


#Adding in currency formatting to revenue_impact column
late_creative_pivot['Revenue_Impact'] = late_creative_pivot['Revenue_Impact'].apply(lambda x: '${:,.2f}'.format(x))


# In[42]:


#Adding in subtotal formatting for the Account_Executive column
late_creative_pivot_final = df2.groupby('Account_Executive').apply(lambda sub: sub.pivot_table(
    index=['Account_Executive', 'Advertiser'],
    values=['Name','Revenue_Impact','Total_Month_Dollars'],
    aggfunc={'Name': pd.Series.nunique, 'Total_Month_Dollars': np.sum, 'Revenue_Impact':np.sum},
    margins=True,
    margins_name= 'SubTotal',
))


# In[43]:


late_creative_pivot_final


# In[44]:


late_creative_pivot_final.loc[('', 'Total', '')] = late_creative_pivot_final.sum()


# In[45]:


#Removing duplicate Account_Executive column
late_creative_pivot_final.index = late_creative_pivot_final.index.droplevel(0)


# In[46]:


late_creative_pivot_final


# In[47]:


#Formatting
late_creative_pivot_final['Total_Month_Dollars']=late_creative_pivot_final['Total_Month_Dollars'].apply(lambda x: '${:,.2f}'.format(x))


# In[48]:


late_creative_pivot_final['Revenue_Impact']=late_creative_pivot_final['Revenue_Impact'].apply(lambda x: '${:,.2f}'.format(x))


# In[49]:


late_creative_pivot_final['Name']=late_creative_pivot_final['Name'].apply(lambda x: '{0:.3g}'.format(x))


# In[50]:


#Adding in notes column 
late_creative_pivot_final['Notes']=""


# In[51]:


late_creative_pivot_final.rename(columns={'Revenue_Impact': 'Pending '+ current_month_text +" Revenue"}, inplace=True)


# In[52]:


late_creative_pivot_final.rename(columns={'Total_Month_Dollars': 'Est. Revenue Impact'}, inplace=True)


# In[53]:


late_creative_pivot_final=late_creative_pivot_final.reset_index()


# In[54]:


late_creative_pivot_final


# In[55]:


#Creating list of account executives; We need to append account executive names to their respective subtotal
list1= list(late_creative_pivot_final['Account_Executive'])
list1


# In[57]:


#Creating function that extracts all "subtotal" strings from list1
def list_duplicates_of(seq,item):
    start_at = -1
    locs = []
    while True:
        try:
            loc = seq.index(item,start_at+1)
        except ValueError:
            break
        else:
            locs.append(loc)
            start_at = loc
    return locs

source = list1
index_subtotal = list_duplicates_of(source, 'SubTotal')


# In[58]:


#Concatenates subtotal string to the above account executive
for i in index_subtotal:
    list1[i]=list1[i-1]+" " +list1[i]


# In[59]:


list1


# In[60]:


late_creative_pivot_final['Account_Executive'] = list1


# In[61]:


late_creative_pivot_final


# In[62]:


#Variable currency conversions


# In[63]:


Late_Creative = df2['Revenue_Impact'].sum()
Late_Creative


# In[64]:


Late_Creative_str = '${:,.0f}'.format(Late_Creative)
Late_Creative_str


# In[65]:


##Calculating Total Underdelivery and formatting into currency


# In[66]:


Current_Month_Estimated_UD = df['Current_Month_Estimated_Underdelivery_Amount'].sum()
Current_Month_Estimated_UD


# In[67]:


Current_Month_Estimated_UD_str = '${:,.0f}'.format(Current_Month_Estimated_UD)
Current_Month_Estimated_UD_str


# In[68]:


Late_Creative_Rev_Impact=df2['Revenue_Impact'].sum()
Late_Creative_Rev_Impact


# In[69]:


Late_Creative_Rev_Impact_str = '${:,.0f}'.format(Late_Creative_Rev_Impact)
Late_Creative_Rev_Impact_str


# In[70]:


Total_Underdelivery = Current_Month_Estimated_UD + Late_Creative_Rev_Impact
Total_Underdelivery


# In[71]:


Total_Underdelivery_str = '${:,.0f}'.format(Total_Underdelivery)
Total_Underdelivery_str


# In[72]:


Active_Revenue_At_Risk=df['Current_Month_Estimated_Underdelivery_Amount'].where((df['Ad_Status']=='active')).sum()
Active_Revenue_At_Risk


# In[73]:


Active_Revenue_At_Risk_str = '${:,.0f}'.format(Active_Revenue_At_Risk)
Active_Revenue_At_Risk_str


# In[74]:


###Creating week-over-week calculations###
#Step 1: Create formatted aggregate variables


# In[75]:


Not_Live = Suspended + Completed + Late_Creative
Not_Live


# In[76]:


Not_Live_str = '${:,.0f}'.format(Not_Live)
Not_Live_str


# In[77]:


Suspended_str = '${:,.0f}'.format(Suspended)
Suspended_str


# In[78]:


Completed_str = '${:,.0f}'.format(Completed)
Completed_str


# In[79]:


#Printing Last File

list_of_files = glob.iglob('/Users/Jeffrey.Lu@ibm.com/Desktop/Pacing_Reports/*.xlsx')
latest_file = max(list_of_files, key=os.path.getctime)
print(latest_file)


# In[80]:


#Importing files from previous Pacing Report and going into calculations summary (May need to change directory)


# In[81]:


os.chdir('/Users/Jeffrey.Lu@ibm.com/Desktop/Pacing_Reports')


# In[82]:


df3=pd.read_excel(latest_file, sheet_name ='Calculations Summary')


# In[83]:


df3.head()


# In[84]:


#Create Calculation for the Previous Pacing Report:

prev_Active_Revenue_At_Risk = df3['Underdelivery Amount'][0]

prev_Active_Not_Live = df3['Underdelivery Amount'][1]

prev_Suspended = df3['Underdelivery Amount'][2]

prev_Completed = df3['Underdelivery Amount'][3]

prev_Late_Creative = df3['Underdelivery Amount'][4]


# In[85]:


#Create a calculations sheet for the current report:

d = {'Underdelivery Breakout': ['Current: Active_Revenue_At_Risk','Current: Active_Not_Live','Current: Suspended',                                'Current: Completed','Current: Late Creative'],'Underdelivery Amount':     [Active_Revenue_At_Risk, Not_Live, Suspended, Completed,Late_Creative_Rev_Impact]}
df_calculations_sheet = pd.DataFrame(data = d)


# In[86]:


df_calculations_sheet


# In[87]:


#Percentage Change Variable Calculations

Percentage_Change_Active_Revenue_At_Risk = ((Active_Revenue_At_Risk - prev_Active_Revenue_At_Risk)-1)*100
Percentage_Change_Active_Revenue_At_Risk_str = str((int(Percentage_Change_Active_Revenue_At_Risk))) + '%'
Percentage_Change_Active_Revenue_At_Risk_str


# In[88]:


Percentage_Active_Not_Live = ((Not_Live - prev_Active_Not_Live)-1)*100
Percentage_Active_Not_Live_str = str((int(Percentage_Change_Active_Revenue_At_Risk))) + '%'
Percentage_Active_Not_Live_str


# In[89]:


#Exporting Dataframes, Pivots, and Calculation Sheets to Excel Report

writer = pd.ExcelWriter(time.strftime('Pacing_Report %m-%d-%Y.xlsx'))

df.to_excel(writer,sheet_name='Pacing')
df_pivot_advertiser.to_excel(writer,sheet_name='Advertiser Breakout')
df2.to_excel(writer,sheet_name='Late Creative')
late_creative_pivot_final.to_excel(writer,sheet_name='LC Pivot')
df_calculations_sheet.to_excel(writer, sheet_name = 'Calculations Summary')


writer.save()


# In[90]:


#Completed List to Word Format Manipulation

Completed_df = df[['Current_Month_Estimated_Underdelivery_Amount','Advertiser']].where((df['Ad_Status']=='complete')&(df['Current_Month_Estimated_Underdelivery_Amount']!=0))

Completed_df = Completed_df.groupby(['Advertiser']).sum()

Completed_df['Current_Month_Estimated_Underdelivery_Amount'] = Completed_df['Current_Month_Estimated_Underdelivery_Amount'].apply(lambda x: '${:,.0f}'.format(x))

Completed_df = Completed_df.reset_index()

Completed_df['Advertiser-Current_Month_UD'] = Completed_df['Advertiser'] +"- "+ Completed_df['Current_Month_Estimated_Underdelivery_Amount'].map(str)

Completed_df


# In[91]:


#Completed List

list_Completed = []

for i in Completed_df['Advertiser-Current_Month_UD']:
    list_Completed.append(i)

if list_Completed == []:
    print('Completed lines are at $0 UD')
else:
    print(list_Completed)


# In[92]:


#Suspended List to Word Format Manipulation

Suspended_df = df[['Current_Month_Estimated_Underdelivery_Amount','Advertiser']].where((df['Ad_Status']=='suspended')&(df['Current_Month_Estimated_Underdelivery_Amount']!=0))

Suspended_df = Suspended_df.groupby(['Advertiser']).sum()

Suspended_df['Current_Month_Estimated_Underdelivery_Amount'] = Suspended_df['Current_Month_Estimated_Underdelivery_Amount'].apply(lambda x: '${:,.0f}'.format(x))

Suspended_df = Suspended_df.reset_index()

Suspended_df['Advertiser-Current_Month_UD'] = Suspended_df['Advertiser'] +"- "+ Suspended_df['Current_Month_Estimated_Underdelivery_Amount'].map(str)

Suspended_df




# In[93]:


#Suspended List

list_Suspended = []

for i in Suspended_df['Advertiser-Current_Month_UD']:
    list_Suspended.append(i)

if list_Suspended == []:
    print('Suspended lines are at $0 UD')
else:
    print(list_Suspended)


# In[94]:


#WFX List to Word Format Manipulation
df1 = df[~df['Product_Name'].str.contains('Preroll', na = False)]                              [df['Ad_Type'].str.contains('STANDARD|HOUSE|AD_EXCHANGE', na = False)]                              .where((df['Ad_Status']=='active')&(df['Product_Category']=='WEATHRFX')&(df['Current_Month_Estimated_Underdelivery_Amount']!=0))


WFX_df = df1[['Current_Month_Estimated_Underdelivery_Amount','Advertiser']]

WFX_df = WFX_df.groupby(['Advertiser']).sum()

WFX_df['Current_Month_Estimated_Underdelivery_Amount'] = WFX_df['Current_Month_Estimated_Underdelivery_Amount'].apply(lambda x: '${:,.0f}'.format(x))

WFX_df = WFX_df.reset_index()

WFX_df['Advertiser-Current_Month_UD'] = WFX_df['Advertiser'] +"- "+ WFX_df['Current_Month_Estimated_Underdelivery_Amount'].map(str)


WFX_df


# In[95]:


#WFX List

list_WFX = []

for i in WFX_df['Advertiser-Current_Month_UD']:
    list_WFX.append(i)

if list_WFX == []:
    print('WFX lines are at $0 UD')
else:
    print(list_WFX)


# In[96]:


#AFX List to Word Format Manipulation
df2 = df[~df['Product_Name'].str.contains('Preroll', na = False)]                              [df['Ad_Type'].str.contains('STANDARD|HOUSE|AD_EXCHANGE', na = False)]                              .where((df['Ad_Status']=='active')&(df['Product_Category']=='AUDIENCEFX')&(df['Current_Month_Estimated_Underdelivery_Amount']!=0))


AFX_df = df2[['Current_Month_Estimated_Underdelivery_Amount','Advertiser']]

AFX_df = AFX_df.groupby(['Advertiser']).sum()

AFX_df['Current_Month_Estimated_Underdelivery_Amount'] = AFX_df['Current_Month_Estimated_Underdelivery_Amount'].apply(lambda x: '${:,.0f}'.format(x))

AFX_df = AFX_df.reset_index()

AFX_df['Advertiser-Current_Month_UD'] = AFX_df['Advertiser'] +"- "+ AFX_df['Current_Month_Estimated_Underdelivery_Amount'].map(str)


AFX_df


# In[97]:


#AFX List

list_AFX = []

for i in AFX_df['Advertiser-Current_Month_UD']:
    list_AFX.append(i)

if list_AFX == []:
    print('WFX lines are at $0 UD')
else:
    print(list_AFX)


# In[98]:


#JFX List to Word Format Manipulation
df3 = df[~df['Product_Name'].str.contains('Preroll', na = False)]                              [df['Ad_Type'].str.contains('STANDARD|HOUSE|AD_EXCHANGE', na = False)]                              .where((df['Ad_Status']=='active')&(df['Product_Category']=='JFX')&(df['Current_Month_Estimated_Underdelivery_Amount']!=0))


JFX_df = df3[['Current_Month_Estimated_Underdelivery_Amount','Advertiser']]

JFX_df = JFX_df.groupby(['Advertiser']).sum()

JFX_df['Current_Month_Estimated_Underdelivery_Amount'] = JFX_df['Current_Month_Estimated_Underdelivery_Amount'].apply(lambda x: '${:,.0f}'.format(x))

JFX_df = JFX_df.reset_index()

JFX_df['Advertiser-Current_Month_UD'] = JFX_df['Advertiser'] +"- "+ JFX_df['Current_Month_Estimated_Underdelivery_Amount'].map(str)


JFX_df


# In[99]:


#JFX List

list_JFX = []

for i in JFX_df['Advertiser-Current_Month_UD']:
    list_JFX.append(i)

if list_JFX == []:
    print('JFX lines are at $0 UD')
else:
    print(list_JFX)


# In[100]:


#Preroll List to Word Format Manipulation

df4 = df[df['Product_Name'].str.contains('Preroll', na = False)]                              [df['Ad_Type'].str.contains('STANDARD|HOUSE|AD_EXCHANGE', na = False)]                              .where((df['Ad_Status']=='active')&(df['Current_Month_Estimated_Underdelivery_Amount']!=0))


Preroll_df = df4[['Current_Month_Estimated_Underdelivery_Amount','Advertiser']]

Preroll_df = Preroll_df.groupby(['Advertiser']).sum()

Preroll_df['Current_Month_Estimated_Underdelivery_Amount'] = Preroll_df['Current_Month_Estimated_Underdelivery_Amount'].apply(lambda x: '${:,.0f}'.format(x))

Preroll_df = Preroll_df.reset_index()

Preroll_df['Advertiser-Current_Month_UD'] = Preroll_df['Advertiser'] +"- "+ Preroll_df['Current_Month_Estimated_Underdelivery_Amount'].map(str)


Preroll_df


# In[101]:


#Preroll List

list_Preroll = []

for i in Preroll_df['Advertiser-Current_Month_UD']:
    list_Preroll.append(i)

if list_Preroll == []:
    print('Preroll lines are at $0 UD')
else:
    print(list_Preroll)


# In[102]:


#Sponsorship List to Word Format Manipulation

df5 = df.where((df['Ad_Status']=='active')&(df['Ad_Type']=='SPONSORSHIP')&(df['Current_Month_Estimated_Underdelivery_Amount']!=0))


Spon_df = df5[['Current_Month_Estimated_Underdelivery_Amount','Advertiser']]

Spon_df = Spon_df.groupby(['Advertiser']).sum()

Spon_df['Current_Month_Estimated_Underdelivery_Amount'] = Spon_df['Current_Month_Estimated_Underdelivery_Amount'].apply(lambda x: '${:,.0f}'.format(x))

Spon_df = Spon_df.reset_index()

Spon_df['Advertiser-Current_Month_UD'] = Spon_df['Advertiser'] +"- "+ Spon_df['Current_Month_Estimated_Underdelivery_Amount'].map(str)


Spon_df


# In[103]:


#Sponsorship List

list_Spon = []

for i in Spon_df['Advertiser-Current_Month_UD']:
    list_Spon.append(i)

if list_Spon == []:
    print('Sponsorship lines are at $0 UD')
else:
    print(list_Spon)


# In[104]:


#Exporting Data to Word Doc


# In[105]:


doc = docx.Document()

paraObject = doc.add_paragraph(current_month_year_text + " Revenue at Risk (Under-Delivery+Late Creative) " + Total_Underdelivery_str)

doc.add_paragraph("Active (Live) - " + Active_Revenue_At_Risk_str + " (" + Percentage_Change_Active_Revenue_At_Risk_str + ")")

doc.add_paragraph('         *Active under-delivering campaigns')

doc.save('Pacing Report Word.docx')


# In[106]:


#All Formatted Lines
line1 = current_month_year_text + " Revenue at Risk (Under-Delivery+Late Creative) " + Total_Underdelivery_str
line2 = "Active (Live) - " + Active_Revenue_At_Risk_str + " (" + Percentage_Change_Active_Revenue_At_Risk_str + ")"
line3 = "Not Live - " + Not_Live_str + " (" + Percentage_Active_Not_Live_str + ")"
 
#Doc Object

doc = docx.Document()

#Line 1 Formatted Bold
p = doc.add_paragraph()

runner = p.add_run(line1)
runner.bold = True

#Extra Space
doc.add_paragraph()

#Line 2 Formatted Bold
p1 = doc.add_paragraph()

runner2 = p1.add_run(line2)
runner2.bold = True

#Line 3
doc.add_paragraph('         * Active under-delivering campaigns')

#Line 4 Formatted Bold
p2 = doc.add_paragraph()

runner3 = p2.add_run(line3)
runner3.bold = True

#Line 5  
doc.add_paragraph('         * Suspended- ' + Suspended_str)

#Line 6
doc.add_paragraph('         * Completed- ' + Completed_str)

#Line 7
doc.add_paragraph('         * Late Creative- ' + Late_Creative_Rev_Impact_str)



doc.save(time.strftime('Pacing Report Word %m-%d-%Y.docx'))


# In[107]:


#Report Email Section


# In[108]:


#Creating Email Variables
email_user = 'Insert_Email_Here@gmail.com'
email_password = 'Enter_Password_Here'
email_send_to = 'Enter_Recipient_Email_Addresses_Here'
email_subject = time.strftime('Pacing_Report Materials %m-%d-%Y')


# In[112]:


#Email Object and Details


# In[113]:


msg = MIMEMultipart()


# In[114]:


msg['From'] = email_user


# In[115]:


msg['To'] = email_send_to


# In[116]:


msg['Subject'] = email_subject


# In[117]:


#Email Body > Attach body to Message Object 


# In[118]:


body = 'Please see the attaced pacing materials below. Have a great day!'


# In[119]:


msg.attach(MIMEText(body,'plain'))


# In[121]:


#File Attachment Document variables, octet-streams and base64 encoding


# In[122]:


filename1 = time.strftime('Pacing Report Word %m-%d-%Y.docx')


# In[123]:


filename2 = time.strftime('Pacing_Report %m-%d-%Y.xlsx')


# In[124]:


attachment1 = open(filename1, 'rb')


# In[125]:


attachment2 = open(filename2, 'rb')


# In[127]:


part1 = MIMEBase('application','octet-stream')


# In[128]:


part2 = MIMEBase('application','octet-stream')


# In[129]:


part1.set_payload((attachment1).read())


# In[130]:


part2.set_payload((attachment2).read())


# In[131]:


encoders.encode_base64(part1)


# In[132]:


encoders.encode_base64(part2)


# In[133]:


part1.add_header('Content-Disposition', "attachment; filename= " + filename1)


# In[134]:


part2.add_header('Content-Disposition', "attachment; filename= " + filename2)


# In[135]:


#Converting all of the Message as plain text


# In[136]:


msg.attach(part1)


# In[137]:


msg.attach(part2)


# In[138]:


text = msg.as_string()


# In[139]:


#Establishing SMTP Connection and tls encryption


# In[140]:


conn = smtplib.SMTP('smtp.gmail.com',587)
type(conn)


# In[141]:


conn


# In[142]:


conn.ehlo()


# In[143]:


conn.starttls()


# In[144]:


conn.login(email_user, email_password)


# In[145]:


conn.sendmail(email_user,email_send_to, text)


# In[146]:


conn.quit()


# In[147]:


###Google Sheets API Connection for AE Update Notes on Campaign Delivery###


# In[151]:


os.chdir('/Users/Jeffrey.Lu@ibm.com/Desktop')


# In[152]:


scope = ['https://spreadsheets.google.com/feeds', 'https://www.googleapis.com/auth/drive']


# In[153]:


creds = ServiceAccountCredentials.from_json_keyfile_name('client_credentials.json', scope)


# In[154]:


client = gspread.authorize(creds)


# In[ ]:


# Opening up the Google Sheets Doc and inserting new sheet


# In[221]:


sheet = client.open('Late Creative Doc')


# In[222]:


sheet.add_worksheet(title= time.strftime('Late Creative %m-%d-%Y'), rows="100", cols="20") 


# In[225]:


wks = sheet.worksheet(time.strftime('Late Creative %m-%d-%Y'))


# In[ ]:


#Google sheets doesn't have insert column command so we need the pivot data laid out as rows, thus the list variables below


# In[188]:


late_creative_pivot_final_ae = list(late_creative_pivot_final['Account_Executive'])


# In[189]:


late_creative_pivot_final_adv= list(late_creative_pivot_final['Advertiser'])


# In[190]:


late_creative_pivot_final_name= list(late_creative_pivot_final['Name'])


# In[235]:


late_creative_pivot_final_pend = list(late_creative_pivot_final['Pending '+ current_month_text +" Revenue"])


# In[192]:


late_creative_pivot_final_imp = list(late_creative_pivot_final['Est. Revenue Impact'])


# In[193]:


late_creative_pivot_final_headers = ['Account_Executive','Advertiser','Name','Pending September Revenue', 'Est. Revenue Impact', 'Notes']


# In[194]:


late_creative_pivot_final


# In[195]:


#Figuring out corresponsindg Transpose formula alphabet letter 


# In[196]:


late_creative_pivot_len = len(late_creative_pivot_final['Account_Executive'])


# In[197]:


alphabet_list = ['a','b','c','d','e','f','g','h','i','j','k','l',
               'm','n','o','p','q','r','s','t','u','v','w','x','y','z']


# In[198]:


alphabet_number=late_creative_pivot_len-1


# In[199]:


alphabet_number


# In[200]:


alphabet_var = alphabet_list[alphabet_number]
alphabet_var


# In[201]:


#Writing Late Creative materials to Google Sheets


# In[226]:


row1 = late_creative_pivot_final_ae
index1 = 1
wks.insert_row(row1, index1)


# In[203]:


row2 = late_creative_pivot_final_adv
index2 = 2
wks.insert_row(row2, index2)


# In[204]:


row3 = late_creative_pivot_final_name
index3 = 3
wks.insert_row(row3, index3)


# In[205]:


row4 = late_creative_pivot_final_pend
index4 = 4
wks.insert_row(row4, index4)


# In[206]:


row5 = late_creative_pivot_final_imp
index5 = 5
wks.insert_row(row5, index5)


# In[207]:


row9 = late_creative_pivot_final_headers
index9 = 9
wks.insert_row(row9, index9)


# In[208]:


wks.update_cell(10,1, '=TRANSPOSE(A1:'+alphabet_var+'5)')


# In[ ]:


###Cron Job Established in Terminal###
# 0 11 * * 1,4 /anaconda3/bin/python /Users/Jeffrey.Lu@ibm.com/Desktop/Pacing_Automation.py


# In[ ]:


###Future Improvement List###
#1.) Automatic data importing process via pyodbc or selenium.
#2.) Better report formatting could be written. Write python code that will insert VBA script into the excel/Google Sheets
#    export and run it.


# In[ ]:


###Modifications that need to be made for when this script is inherited###

#1.)Script is customized for this mac computer. If changed, we need to change Current Work Directory for Pacing report retrival

#2.)1st time script will run, it will use the current excel output as the current output and the previous output until
#  the program is run a second time for the next report where it can then use the last outputted report as the previous report
#  for week-over-week calculations

#3.) Email SMPT credentials username and password will need to be changed

#4.) Cron Job script will need to be entered and the variables will need to change

#5.) Google Spreadsheet and Drive API Credentails need to change

#6.) Google API Access to Drive and Sheets will need to be granted 

#7.) All packages will need to be downloaded via pip or sudo pip

